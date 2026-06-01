# -*- coding: utf-8 -*-
"""Build the Chinese DOCX report for the current force-control simulation."""

from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = PROJECT_ROOT / "reports"
OUTPUT_DOCX = REPORT_DIR / "力控螺旋搜索仿真项目总结.docx"

BODY_FONT = "Microsoft YaHei"
CODE_FONT = "Consolas"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"


def read_metrics() -> dict[str, float]:
    path = REPORT_DIR / "force_control_summary.txt"
    metrics: dict[str, float] = {}
    if not path.exists():
        return metrics
    for line in path.read_text(encoding="utf-8").splitlines():
        if ":" not in line or line.startswith("-"):
            continue
        key, value = line.split(":", 1)
        try:
            metrics[key.strip()] = float(value.strip())
        except ValueError:
            continue
    return metrics


def read_phase_summary() -> list[tuple[str, float, float, float, float]]:
    path = REPORT_DIR / "force_control_metrics.csv"
    if not path.exists():
        return []
    rows = list(csv.DictReader(path.open("r", encoding="utf-8")))
    groups: dict[str, list[dict[str, str]]] = {}
    for row in rows:
        phase = row["phase"]
        groups.setdefault(phase, []).append(row)
    selected = []
    for phase in [
        "align gripper center to object center using joint1/joint2 only",
        "descend to grasp height using joint5 only",
        "Archimedean spiral search using joint1/joint2 only",
        "final insertion using joint5 with compensated joint6 screwing",
    ]:
        data = groups.get(phase)
        if not data:
            continue
        t0 = float(data[0]["time"])
        t1 = float(data[-1]["time"])
        max_force = max(float(r["contact_force_n"]) for r in data)
        xy_final = float(data[-1]["peg_hole_xy_error_m"])
        depth_final = float(data[-1]["inserted_depth_m"])
        selected.append((phase, t1 - t0, max_force, xy_final, depth_final))
    return selected


def set_font(run, name: str = BODY_FONT, size: float | None = None, bold: bool | None = None, color: str | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT, size: float = 9.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("轴孔装配力控螺旋搜索仿真项目总结")
    set_font(r, size=22, bold=True, color="17365D")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("基于 JTC/TCC/FSC 的 6DOF 机械臂力控、螺旋搜索与补偿式 Screwing 插入验证")
    set_font(r, size=11, color="555555")


def add_callout(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, CALLOUT)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title + "：")
    set_font(r, size=10.5, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_font(r, size=10.5)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.167
        r = p.add_run(item)
        set_font(r, size=10.5)


def add_formula(doc: Document, image_name: str, caption: str) -> None:
    path = REPORT_DIR / image_name
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(5.6))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_font(r, size=9, color="666666")
    r.italic = True


def add_figure(doc: Document, image_name: str, caption: str, width: float = 6.25) -> None:
    path = REPORT_DIR / image_name
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_font(r, size=9, color="555555")
    r.italic = True


def add_metrics_table(doc: Document, metrics: dict[str, float]) -> None:
    labels = [
        ("总仿真时长", "duration_s", "s"),
        ("最大瞬时接触力", "max_contact_force_n", "N"),
        ("平均接触力", "mean_contact_force_n", "N"),
        ("最终插入深度", "final_inserted_depth_m", "m"),
        ("插入阶段深度变化", "insertion_phase_depth_change_m", "m"),
        ("最终孔-轴平面误差", "final_xy_error_m", "m"),
        ("Joint3 最大偏移", "max_abs_joint3_rad", "rad"),
        ("Joint4 最大偏移", "max_abs_joint4_rad", "rad"),
        ("Joint6 最终回转圈数", "final_joint6_turns", "turns"),
        ("去除 Joint6 后最大跟踪误差", "max_tracking_error_without_q6", ""),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ["指标", "数值", "含义"]):
        set_cell_text(cell, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, LIGHT_GRAY)
    for label, key, unit in labels:
        if key not in metrics:
            continue
        value = metrics[key]
        row = table.add_row().cells
        set_cell_text(row[0], label)
        set_cell_text(row[1], f"{value:.5g} {unit}".strip(), align=WD_ALIGN_PARAGRAPH.CENTER)
        meaning = {
            "duration_s": "包含抓取、移动、螺旋搜索、补偿式 screwing 与保持阶段。",
            "max_contact_force_n": "接触峰值主要来自闭合夹爪和孔口接触瞬间。",
            "mean_contact_force_n": "反映力控过程的整体接触强度。",
            "final_inserted_depth_m": "圆柱体最终进入孔深内的深度。",
            "insertion_phase_depth_change_m": "Joint5 下探阶段带来的有效插入量。",
            "final_xy_error_m": "最终圆柱体中心与孔中心的平面距离。",
            "max_abs_joint3_rad": "用于量化 Link3 姿态保持效果。",
            "max_abs_joint4_rad": "用于量化 Link4 竖直稳定性。",
            "final_joint6_turns": "补偿式 screwing 保留的末端回转圈数。",
            "max_tracking_error_without_q6": "排除多圈 joint6 后的 arm 跟踪误差。",
        }[key]
        set_cell_text(row[2], meaning, size=9.0)


def add_phase_table(doc: Document, phase_rows: list[tuple[str, float, float, float, float]]) -> None:
    if not phase_rows:
        return
    names = {
        "align gripper center to object center using joint1/joint2 only": "抓取前平面对准",
        "descend to grasp height using joint5 only": "Joint5 下探抓取",
        "Archimedean spiral search using joint1/joint2 only": "阿基米德螺旋搜索",
        "final insertion using joint5 with compensated joint6 screwing": "补偿式 screwing 插入",
    }
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["阶段", "持续时间/s", "最大接触力/N", "末端XY误差/mm", "插入深度/mm"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_text(cell, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
        set_cell_shading(cell, LIGHT_GRAY)
    for phase, dt, force, xy, depth in phase_rows:
        cells = table.add_row().cells
        values = [names.get(phase, phase), f"{dt:.3f}", f"{force:.2f}", f"{xy * 1000:.2f}", f"{depth * 1000:.2f}"]
        for i, value in enumerate(values):
            set_cell_text(cells[i], value, align=WD_ALIGN_PARAGRAPH.CENTER if i else WD_ALIGN_PARAGRAPH.LEFT, size=8.7)


def build_doc() -> None:
    REPORT_DIR.mkdir(exist_ok=True)
    metrics = read_metrics()
    phase_rows = read_phase_summary()

    doc = Document()
    style_document(doc)
    add_title(doc)
    add_callout(
        doc,
        "摘要",
        "当前版本将底层控制从刚性位置控制推进到基于 MuJoCo motor actuator 的力/力矩控制，并在 JTC/TCC/FSC 框架下完成抓取、螺旋搜索、Joint5 插入和 Joint6 补偿式 screwing。新的控制逻辑允许接触反力进入闭环，能够用接触力、跟踪误差、插入深度和关节稳定性量化装配过程。",
    )

    doc.add_heading("1. 控制框架更新与优势", level=1)
    doc.add_paragraph(
        "项目当前采用的底层控制不再直接写入关节位置，而是向 MuJoCo 的 motor actuator 写入广义力。对于转动关节，控制量是力矩；对于 Joint5、夹爪等平动关节，控制量是直线力。这样做的核心收益是：接触不会被刚性位置命令覆盖，机械臂在孔口、圆柱体和夹爪之间产生的反力可以反映到关节误差和控制输入中。"
    )
    add_bullets(
        doc,
        [
            "JTC（Jacobian Transpose Controller）将任务空间力/力矩映射为关节广义力，适合生成接触力和插入力。",
            "TCC（Task-oriented Coordinate Controller）在末端操作空间中控制位置和姿态，使螺旋搜索轨迹不依赖单纯关节空间插值。",
            "FSC（Force-motion Separating Controller）将插入方向的力生成与平面位置控制分离：x/y 和姿态保持反馈控制，z 方向下压力以 feed-forward 方式施加。",
            "PID/阻抗项允许外力造成有限位姿误差，并通过积分与阻尼消除稳态误差、限制振动。",
        ],
    )

    doc.add_heading("2. 不同阶段的力控制公式", level=1)
    doc.add_paragraph("控制器的总力矩结构参考论文中的 JTC/TCC/FSC 写法，同时加入 MuJoCo bias 补偿、关节姿态保持和库仑摩擦补偿。公式图片由 LaTeX/mathtext 渲染后嵌入，便于在 Word/PDF 中稳定显示。")
    add_formula(doc, "formula_total_tau.png", "式 1  总控制律：任务空间 wrench 经 J^T 映射到关节广义力。")
    add_formula(doc, "formula_task_pid.png", "式 2  TCC：任务空间 PID 根据 xd - xc 生成反馈力和力矩。")
    add_formula(doc, "formula_fsc.png", "式 3  FSC：运动轴使用反馈 wrench，受力轴使用 feed-forward wrench。")
    add_formula(doc, "formula_spiral.png", "式 4  螺旋搜索：圆柱体中心沿阿基米德螺旋在固定平面内搜索孔口。")
    add_formula(doc, "formula_screw_comp.png", "式 5  补偿式 screwing：Joint6 自转时由 Joint1/2 抵消偏心项，使 peg 中心保持在孔中心附近。")

    doc.add_heading("3. 阶段化控制逻辑", level=1)
    add_phase_table(doc, phase_rows)
    doc.add_paragraph(
        "表中可以看到，螺旋搜索阶段主要消耗时间但接触力较低；插入阶段通过 Joint5 下探和 Joint6 约两圈回转完成最终装配。由于 Joint6 轴线并不穿过圆柱体中心，当前代码在 screwing 阶段引入 Joint1/Joint2 的补偿反解，避免圆柱体中心绕偏置轴形成非物理圆周运动。"
    )

    doc.add_heading("4. 量化仿真结果", level=1)
    add_metrics_table(doc, metrics)

    doc.add_heading("5. 图表分析", level=1)
    add_figure(doc, "contact_force_and_depth.png", "图 1  接触力与插入深度。插入深度在最终阶段逐步增加，接触峰值主要来自接触建立和孔口约束。")
    add_figure(doc, "joint_stability.png", "图 2  Joint3/Joint4 稳定性与 Joint6 回转圈数。Joint6 保留 screwing，Joint3/4 偏移被姿态保持项限制。")
    add_figure(doc, "tracking_error.png", "图 3  关节跟踪误差分解。完整误差包含 Joint6 多圈回转，去除 Joint6 后可观察主机械臂平动/姿态跟踪质量。")
    add_figure(doc, "peg_xy_path.png", "图 4  圆柱体中心平面轨迹。轨迹体现先偏置搜索、后靠近孔中心并进行补偿式插入。", width=5.6)

    doc.add_heading("6. 对机械臂控制带来的变化", level=1)
    add_bullets(
        doc,
        [
            "从“位置必须到达”转为“在力矩限幅内趋近目标”，因此碰撞会产生可观测的跟踪误差和接触力峰值。",
            "螺旋搜索阶段可以显式用 f = Kp(xd - xc) 描述平面位置反馈，目标轨迹和实际接触响应分离得更清楚。",
            "插入阶段 z 方向由 FSC 提供向下 feed-forward 力，x/y 仍由位置反馈维持孔轴对准，降低了下探时横向穿模和抖动风险。",
            "补偿式 screwing 保留了 Joint6 自转的装配动作，同时通过 Joint1/Joint2 纠正由结构偏置带来的圆柱体中心公转。",
        ],
    )

    doc.add_heading("7. 创新点", level=1)
    add_bullets(
        doc,
        [
            "将论文中的 JTC/TCC/FSC 思路落到 MuJoCo motor actuator 层，实现关节力/力矩输入而非直接 qpos 覆盖。",
            "把轴孔装配拆分为抓取、恒高接触、阿基米德螺旋搜索、FSC 下探和补偿式 screwing 五个阶段，便于单独调参与量化。",
            "针对当前 v7 机械臂 Joint6 轴线与 peg 中心不共线的问题，引入 screwing 偏心补偿逻辑，而不是简单关闭 Joint6。",
            "用力、位姿误差、插入深度、关节稳定性和 peg 平面轨迹构成评价指标，不只依靠 GUI 观察判断效果。",
        ],
    )

    doc.add_heading("8. 不足与局限", level=1)
    add_bullets(
        doc,
        [
            "当前仍使用 grasp_lock 稳定夹持，说明纯 STL 接触和摩擦还不足以完全替代真实夹爪物理夹持。",
            "MuJoCo 的软接触模型难以完全复现论文中“失稳后三点接触再 wiggling”的真实接触链，因此项目重点停留在螺旋搜索与插入阶段。",
            "FSC 中的下压力仍主要依赖经验参数，尚未接入力传感器闭环或根据接触状态自适应调节。",
            "Joint6 screwing 的偏心补偿是模型几何下的工程修正；真实机械臂部署前仍需要末端工具坐标系和夹持中心的精确标定。",
            "图表目前基于单 seed 的代表性数据，后续应批量运行多随机初始位置，统计成功率、平均插入时间和接触力分布。",
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("数据来源：scripts/generate_force_control_report.py --seed 1")
    set_font(r, size=9, color="777777")

    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build_doc()
